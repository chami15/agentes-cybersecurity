from langgraph.graph import StateGraph, END
from nodes import scanner_node, remediation_node, rerank_node, node_reporter, triage_node
from state import AgentState
from docx import Document
from docx.shared import Pt

def save_docx(texto_relatorio, nome_arquivo = 'Relatorio_final_pentest.docx'):
    doc = Document()

    titulo = doc.add_heading('Relatorio Vulnerabilidade - Triage AI', 0)

    for linha in texto_relatorio.split('\n'):
        linha = linha.strip()
        if not linha: 
            continue

        if linha.startswith('###'):
            doc.add_heading(linha.replace('#', '').strip(), level=2)

        elif linha.startswith('##'):
            doc.add_heading(linha.replace('#', ''), level=1)

        elif linha.startswith('**') and linha.endswith('**'):
            p = doc.add_paragraph()

            run = p.add_run(linha.replace('*', ''))

            run.bold = True

        else: 
            doc.add_paragraph(linha)

    doc.save(nome_arquivo)
    print(f'Arquivo salvo com sucesso: {nome_arquivo}')

def route_check(state: AgentState):
    scan_data = state['scan_results'][0]

    if 'CLOSED' in scan_data:
        return 'processo finalizado'
    return 'continue_triage'

workflow = StateGraph(AgentState)
workflow.add_node('scanner', scanner_node)
workflow.add_node('triage', triage_node)
workflow.add_node('reranker', rerank_node)
workflow.add_node('remediation', remediation_node)
workflow.add_node('reporter', node_reporter)

workflow.set_entry_point('scanner')

workflow.add_conditional_edges(
    'scanner',
    route_check,
    {
        'end_process': END,
        'continue_triage' : 'triage'
    }
)

workflow.add_edge('triage', 'reranker')
workflow.add_edge('reranker', 'remediation')
workflow.add_edge('remediation', 'reporter')
workflow.add_edge('reporter', END)

app = workflow.compile()

if __name__ == '__main__':
    inputs = {
        'target_ip' : '127.0.0.1',
        'target_port' : 3306,
        'scan_results' : [],
        'verified_vul' : [],
        'messages' : []
    }

    print('Inicializando Vulnerability Triage AI...\n')

    result = app.invoke(inputs)
    report_content = result['final_report']

    print('Relatorio final gerado (visualização)')
    print('------------------------------------')
    print(report_content)
    print('------------------------------------')
    print('Gerando arquivo docx...')
    save_docx(report_content)
    print('processo finalizado!')